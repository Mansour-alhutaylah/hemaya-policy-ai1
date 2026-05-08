"""
routers/reports_export.py

POST /api/reports/export — Compile and stream a full Compliance Package DOCX.

Sections:
  1. Cover page
  2. Executive Summary (compliance scores + AI insights)
  3. Open Gaps table
  4. AI Remediation Drafts with control-satisfaction mapping
  5. Policy Version History
  6. Legal disclaimer
"""
import io
import json
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.security import OAuth2PasswordBearer
from jose import JWTError, jwt
from pydantic import BaseModel
from sqlalchemy import text as sql_text
from sqlalchemy.orm import Session

from backend import auth, models
from backend.database import get_db, set_user_context

router = APIRouter(prefix="/api/reports", tags=["reports-export"])

# ── Auth ──────────────────────────────────────────────────────────────────────

_oauth2 = OAuth2PasswordBearer(tokenUrl="/api/auth/login")


def _get_current_user(
    token: str = Depends(_oauth2),
    db: Session = Depends(get_db),
) -> models.User:
    exc = HTTPException(status_code=401, detail="Could not validate credentials")
    try:
        payload = jwt.decode(token, auth.SECRET_KEY, algorithms=[auth.ALGORITHM])
        email: str = payload.get("sub")
        if not email:
            raise exc
    except JWTError:
        raise exc
    user = db.query(models.User).filter(models.User.email == email).first()
    if not user:
        raise exc
    set_user_context(db, user.id)
    return user


# ── Pydantic ──────────────────────────────────────────────────────────────────

class ExportRequest(BaseModel):
    policy_id: str
    include_draft_text: bool = True


# ── Colour palette (matches the app's Tailwind theme) ────────────────────────

_EMERALD  = RGBColor(16, 185, 129)
_SLATE900 = RGBColor(15,  23,  42)
_SLATE700 = RGBColor(51,  65,  85)
_SLATE500 = RGBColor(100, 116, 139)
_WHITE    = RGBColor(255, 255, 255)
_RED      = RGBColor(220,  38,  38)
_AMBER    = RGBColor(217, 119,   6)
_GREEN    = RGBColor(22, 163,  74)

_SEV_HEX = {
    "Critical": "DC2626",
    "High":     "EA580C",
    "Medium":   "D97706",
    "Low":      "64748B",
}


# ── Document helpers ──────────────────────────────────────────────────────────

def _cell_shading(cell, hex_color: str) -> None:
    """Apply a solid background fill to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    color = _EMERALD if level == 1 else _SLATE900
    size  = Pt(16) if level == 1 else Pt(13)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.size      = size


def _para(
    doc: Document,
    text: str,
    size: int = 10,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor = None,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _header_row(table, *labels: str) -> None:
    """Style the first row as a dark-slate header."""
    row = table.rows[0]
    for i, label in enumerate(labels):
        cell = row.cells[i]
        cell.text = label
        _cell_shading(cell, "0F172A")         # slate-900
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold      = True
                run.font.size      = Pt(9)
                run.font.color.rgb = _WHITE


def _parse_json(val) -> list:
    if val is None:
        return []
    if isinstance(val, (list, dict)):
        return val
    try:
        return json.loads(val)
    except (ValueError, TypeError):
        return []


# ── Data layer ────────────────────────────────────────────────────────────────

def _fetch(db: Session, policy_id: str) -> dict:
    policy = db.query(models.Policy).filter(models.Policy.id == policy_id).first()
    if not policy:
        return {}

    # Latest compliance result per framework
    cr_rows = db.execute(sql_text("""
        SELECT f.name, cr.compliance_score, cr.controls_covered,
               cr.controls_partial, cr.controls_missing, cr.analyzed_at
        FROM   compliance_results cr
        LEFT JOIN frameworks f ON cr.framework_id = f.id
        WHERE  cr.policy_id = :pid
        ORDER  BY cr.analyzed_at DESC
    """), {"pid": policy_id}).fetchall()
    seen, compliance = set(), []
    for r in cr_rows:
        if r[0] not in seen:
            seen.add(r[0])
            compliance.append(r)

    gaps = db.execute(sql_text("""
        SELECT g.control_name, f.name, cl.control_code,
               g.severity, g.status, g.description, g.remediation
        FROM   gaps g
        LEFT JOIN frameworks     f  ON g.framework_id = f.id
        LEFT JOIN control_library cl ON g.control_id  = cl.id
        WHERE  g.policy_id = :pid
        ORDER  BY
            CASE g.severity
                WHEN 'Critical' THEN 1 WHEN 'High' THEN 2
                WHEN 'Medium'   THEN 3 ELSE 4 END,
            g.status
    """), {"pid": policy_id}).fetchall()

    drafts = db.execute(sql_text("""
        SELECT rd.id, cl.control_code, cl.title, f.name,
               rd.remediation_status, rd.missing_requirements,
               rd.suggested_policy_text, rd.section_headers,
               rd.ai_rationale, rd.created_at
        FROM   remediation_drafts rd
        LEFT JOIN control_library cl ON rd.control_id   = cl.id::text
        LEFT JOIN frameworks      f  ON rd.framework_id = f.id
        WHERE  rd.policy_id = :pid
        ORDER  BY rd.created_at DESC
    """), {"pid": policy_id}).fetchall()

    versions = db.execute(sql_text("""
        SELECT version_number, version_type, change_summary, created_at
        FROM   policy_versions
        WHERE  policy_id = :pid
        ORDER  BY version_number
    """), {"pid": policy_id}).fetchall()

    insights = db.execute(sql_text("""
        SELECT title, description, priority, insight_type
        FROM   ai_insights
        WHERE  policy_id = :pid
        ORDER  BY CASE priority
            WHEN 'Critical' THEN 1 WHEN 'High' THEN 2
            WHEN 'Medium'   THEN 3 ELSE 4 END
        LIMIT 8
    """), {"pid": policy_id}).fetchall()

    return {
        "policy":     policy,
        "compliance": compliance,
        "gaps":       gaps,
        "drafts":     drafts,
        "versions":   versions,
        "insights":   insights,
    }


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _build_docx(
    data: dict,
    include_draft_text: bool,
    requester_name: str,
) -> io.BytesIO:
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    policy   = data["policy"]
    pol_name = getattr(policy, "file_name", "Unknown Policy")
    now_str  = datetime.now(timezone.utc).strftime("%B %d, %Y")

    # ── 0. Cover page ─────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("COMPLIANCE PACKAGE")
    r.font.size      = Pt(28)
    r.font.bold      = True
    r.font.color.rgb = _EMERALD

    doc.add_paragraph()

    pn = doc.add_paragraph()
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn = pn.add_run(pol_name)
    rn.font.size      = Pt(16)
    rn.font.color.rgb = _SLATE900

    dn = doc.add_paragraph()
    dn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rd = dn.add_run(now_str)
    rd.font.size      = Pt(11)
    rd.font.color.rgb = _SLATE500

    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp = pp.add_run("Powered by Hemaya AI Compliance")
    rp.font.size      = Pt(9)
    rp.font.italic    = True
    rp.font.color.rgb = _SLATE500

    if requester_name:
        gb = doc.add_paragraph()
        gb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gb.add_run(f"Generated by: {requester_name}").font.size = Pt(9)

    doc.add_page_break()

    # ── 1. Executive Summary ──────────────────────────────────────────────────
    _heading(doc, "1. Executive Summary")

    compliance = data.get("compliance", [])
    if compliance:
        scores = [float(r[1]) for r in compliance if r[1] is not None]
        avg    = round(sum(scores) / len(scores), 1) if scores else 0.0
        _para(doc, f"Overall Compliance Score (average): {avg}%", size=12, bold=True)
        doc.add_paragraph()

        tbl = doc.add_table(rows=1 + len(compliance), cols=5)
        tbl.style = "Table Grid"
        _header_row(tbl, "Framework", "Score", "Covered", "Partial", "Missing")
        for i, r in enumerate(compliance, 1):
            row = tbl.rows[i]
            row.cells[0].text = r[0] or "Unknown"
            row.cells[1].text = f"{r[1]:.1f}%" if r[1] is not None else "—"
            row.cells[2].text = str(r[2] or 0)
            row.cells[3].text = str(r[3] or 0)
            row.cells[4].text = str(r[4] or 0)
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
    else:
        _para(doc, "No compliance analysis found.", italic=True, color=_SLATE500)

    # AI Insights
    insights = data.get("insights", [])
    if insights:
        doc.add_paragraph()
        _heading(doc, "Key AI Insights", level=2)
        for ins in insights:
            p    = doc.add_paragraph(style="List Bullet")
            head = p.add_run(f"[{ins[2]}] {ins[0]}: ")
            head.font.bold = True
            head.font.size = Pt(9)
            body = p.add_run((ins[1] or "")[:300])
            body.font.size = Pt(9)

    doc.add_page_break()

    # ── 2. Open Gaps ──────────────────────────────────────────────────────────
    _heading(doc, "2. Open Gaps & Risks")
    gaps      = data.get("gaps", [])
    open_gaps = [g for g in gaps if (g[4] or "Open") not in ("Resolved", "Closed", "In Progress")]

    if open_gaps:
        _para(doc, f"{len(open_gaps)} open gap(s) requiring attention.", size=9, color=_SLATE500)
        doc.add_paragraph()

        tbl = doc.add_table(rows=1 + len(open_gaps), cols=5)
        tbl.style = "Table Grid"
        _header_row(tbl, "Control", "Framework", "Severity", "Status", "Description")

        for i, g in enumerate(open_gaps, 1):
            row = tbl.rows[i]
            row.cells[0].text = (g[2] or g[0] or "—")
            row.cells[1].text = (g[1] or "—")
            row.cells[2].text = (g[3] or "Medium")
            row.cells[3].text = (g[4] or "Open")
            row.cells[4].text = (g[5] or "")[:180]

            hex_sev = _SEV_HEX.get(g[3] or "Medium", "64748B")
            _cell_shading(row.cells[2], hex_sev)
            for p in row.cells[2].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = _WHITE
                    run.font.bold      = True

            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
    else:
        _para(doc, "No open gaps. All controls are currently addressed.", italic=True, color=_SLATE500)

    doc.add_page_break()

    # ── 3. AI Remediation Drafts ──────────────────────────────────────────────
    _heading(doc, "3. AI Remediation Drafts")
    drafts = data.get("drafts", [])

    if drafts:
        _para(doc, f"{len(drafts)} AI-generated draft(s) on record.", size=9, color=_SLATE500)
        doc.add_paragraph()

        for draft in drafts:
            ctrl_code  = draft[1] or "Unknown Control"
            ctrl_title = draft[2] or ""
            fw_name    = draft[3] or "Unknown Framework"
            status     = (draft[4] or "draft").replace("_", " ").title()
            missing    = _parse_json(draft[5])
            full_text  = draft[6] or ""
            headers    = _parse_json(draft[7])
            rationale  = draft[8] or ""
            created    = draft[9]

            _heading(doc, f"{ctrl_code} — {fw_name}", level=2)
            _para(doc, f"Status: {status}  ·  {ctrl_title}", size=9, color=_SLATE500)
            if created:
                _para(doc, f"Generated: {created.strftime('%Y-%m-%d %H:%M UTC')}", size=9, color=_SLATE500)
            doc.add_paragraph()

            # Why this was suggested (AI rationale)
            if rationale:
                _para(doc, "Why This Was Suggested:", size=10, bold=True)
                _para(doc, rationale, size=9, color=_SLATE700)
                doc.add_paragraph()

            # Missing requirements addressed
            if missing:
                _para(doc, f"Requirements Addressed ({len(missing)}):", size=10, bold=True)
                for req in missing:
                    p = doc.add_paragraph(style="List Number")
                    p.add_run(req).font.size = Pt(9)
                doc.add_paragraph()

            # Control satisfaction mapping
            # Each generated section header is mapped to the triggering control.
            if headers:
                _para(doc, "Control Satisfaction Mapping:", size=10, bold=True)
                for h in headers:
                    p = doc.add_paragraph(style="List Bullet")
                    tick = p.add_run("✓  ")
                    tick.font.bold      = True
                    tick.font.color.rgb = _GREEN
                    tick.font.size      = Pt(9)
                    body = p.add_run(f'"{h}"  satisfies  {ctrl_code}  ({fw_name})')
                    body.font.size = Pt(9)
                doc.add_paragraph()

            # Full draft text (optional, controlled by caller)
            if include_draft_text and full_text:
                _para(doc, "Suggested Policy Addition:", size=10, bold=True)
                txt_p = doc.add_paragraph()
                txt_r = txt_p.add_run(full_text[:3000])
                txt_r.font.size      = Pt(8)
                txt_r.font.color.rgb = _SLATE700
                if len(full_text) > 3000:
                    trunc = doc.add_paragraph()
                    tr    = trunc.add_run("[ Text truncated — view the full draft in the Hemaya platform ]")
                    tr.font.size   = Pt(8)
                    tr.font.italic = True
                doc.add_paragraph()

            # Inter-draft separator
            sep = doc.add_paragraph("─" * 72)
            sep.runs[0].font.size      = Pt(7)
            sep.runs[0].font.color.rgb = _SLATE500
    else:
        _para(doc, "No remediation drafts have been generated yet.", italic=True, color=_SLATE500)
        _para(
            doc,
            "Go to Mapping Review and click 'Accept & Generate Draft' on any flagged mapping.",
            size=9,
            color=_SLATE500,
        )

    doc.add_page_break()

    # ── 4. Policy Version History ─────────────────────────────────────────────
    _heading(doc, "4. Policy Version History")
    versions = data.get("versions", [])

    if versions:
        tbl = doc.add_table(rows=1 + len(versions), cols=4)
        tbl.style = "Table Grid"
        _header_row(tbl, "Version", "Type", "Date", "Change Summary")
        for i, v in enumerate(versions, 1):
            row = tbl.rows[i]
            row.cells[0].text = str(v[0])
            row.cells[1].text = (v[1] or "").replace("_", " ").title()
            ts = v[3]
            row.cells[2].text = ts.strftime("%Y-%m-%d") if ts else "—"
            row.cells[3].text = (v[2] or "")[:160]
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
    else:
        _para(doc, "No version history recorded for this policy.", italic=True, color=_SLATE500)

    doc.add_page_break()

    # ── Disclaimer ────────────────────────────────────────────────────────────
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = disc.add_run(
        "This document was generated automatically by Hemaya AI Compliance. "
        "All AI-generated content is advisory only and must be reviewed by a "
        "qualified compliance officer before being incorporated into official policy documents."
    )
    dr.font.size      = Pt(8)
    dr.font.italic    = True
    dr.font.color.rgb = _SLATE500

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Endpoint ──────────────────────────────────────────────────────────────────

@router.post("/export")
def export_compliance_package(
    body: ExportRequest,
    db: Session = Depends(get_db),
    current_user: models.User = Depends(_get_current_user),
):
    """
    Compile and stream a full Compliance Package DOCX for a given policy.

    Includes executive summary, open gaps, AI remediation drafts with
    control-satisfaction mapping, and policy version history.
    """
    data = _fetch(db, body.policy_id)
    if not data:
        raise HTTPException(status_code=404, detail="Policy not found.")

    requester = (
        f"{current_user.first_name or ''} {current_user.last_name or ''}".strip()
        or current_user.email
    )
    buf = _build_docx(data, body.include_draft_text, requester)

    safe_name = (
        getattr(data["policy"], "file_name", "policy")
        .replace(" ", "_")
        .replace("/", "-")[:60]
    )
    filename = (
        f"himaya_compliance_package_{safe_name}"
        f"_{datetime.now(timezone.utc).strftime('%Y%m%d')}.docx"
    )

    return StreamingResponse(
        buf,
        media_type=(
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
